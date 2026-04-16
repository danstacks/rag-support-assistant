import os
import re
import asyncio
import time
from typing import List, Optional, Set, Dict, Callable
from urllib.parse import urljoin, urlparse, parse_qs
from dataclasses import dataclass, field
import aiohttp
from bs4 import BeautifulSoup
from langchain_core.documents import Document

from app.config import get_settings


@dataclass
class ScrapeConfig:
    """Configuration for web scraping"""
    url: str
    recursive: bool = True
    max_depth: int = 3
    max_pages: int = 500
    allowed_domains: List[str] = field(default_factory=list)
    url_patterns: List[str] = field(default_factory=list)  # Regex patterns to include
    exclude_patterns: List[str] = field(default_factory=list)  # Regex patterns to exclude
    rate_limit: float = 0.5  # Seconds between requests
    headers: Dict[str, str] = field(default_factory=dict)
    auth_token: Optional[str] = None  # Bearer token for authenticated APIs
    cookies: Dict[str, str] = field(default_factory=dict)
    platform: str = "auto"  # auto, confluence, sharepoint, gitbook, docusaurus, mkdocs, sphinx
    # Authentication options
    basic_auth_username: Optional[str] = None
    basic_auth_password: Optional[str] = None
    cookie_string: Optional[str] = None  # Raw cookie string from browser (will be parsed)


def parse_cookie_string(cookie_string: str) -> Dict[str, str]:
    """Parse a cookie string from browser dev tools into a dict.
    
    Accepts formats like:
    - 'name1=value1; name2=value2'
    - 'name1=value1\nname2=value2'
    - Cookie header format from browser
    """
    cookies = {}
    if not cookie_string:
        return cookies
    
    # Normalize separators
    cookie_string = cookie_string.replace('\r\n', ';').replace('\n', ';')
    
    for part in cookie_string.split(';'):
        part = part.strip()
        if '=' in part:
            # Handle cookies with = in value
            key, _, value = part.partition('=')
            key = key.strip()
            value = value.strip()
            if key:  # Skip empty keys
                cookies[key] = value
    
    return cookies


# Platform-specific content selectors
PLATFORM_SELECTORS = {
    "confluence": {
        "content": ["#main-content", ".wiki-content", "[data-testid='content']", ".content-body"],
        "remove": [".page-metadata", ".page-restrictions", "#comments-section", ".confluence-information-macro"],
        "title": [".page-title", "h1"],
    },
    "sharepoint": {
        "content": ["[data-automation-id='pageContent']", ".CanvasSection", "#contentBox", ".ms-rtestate-field"],
        "remove": [".ms-webpart-chrome-title", "[data-automation-id='pageHeader']"],
        "title": ["[data-automation-id='pageTitle']", "h1"],
    },
    "gitbook": {
        "content": ["main", ".markdown-body", "[data-testid='page.contentEditor']", ".page-body"],
        "remove": [".page-toc", ".page-nav", ".gitbook-root > aside"],
        "title": ["h1", ".page-title"],
    },
    "docusaurus": {
        "content": ["article", ".markdown", ".docMainContainer", "main"],
        "remove": [".pagination-nav", ".theme-doc-toc-mobile", ".theme-doc-breadcrumbs"],
        "title": ["h1", ".docTitle"],
    },
    "mkdocs": {
        "content": ["article", ".md-content", "[role='main']", "main"],
        "remove": [".md-sidebar", ".md-footer", ".md-header"],
        "title": ["h1", ".md-content h1"],
    },
    "sphinx": {
        "content": [".document", ".body", "[role='main']", ".rst-content", ".wy-nav-content", "main", "article"],
        "remove": [".sphinxsidebar", ".related", ".footer", ".wy-nav-side", ".rst-versions", ".wy-breadcrumbs"],
        "title": ["h1"],
    },
    "generic": {
        "content": ["main", "article", ".content", ".main-content", "#content", "[role='main']"],
        "remove": ["nav", "footer", "header", "aside", ".sidebar", ".navigation", ".menu"],
        "title": ["h1", "title"],
    }
}


class DocumentLoader:
    def __init__(self):
        self.settings = get_settings()
        self.visited_urls: Set[str] = set()
        self.scrape_stats = {"processed": 0, "skipped": 0, "errors": 0}
    
    async def fetch_url(
        self, 
        session: aiohttp.ClientSession, 
        url: str,
        headers: Optional[Dict[str, str]] = None,
        cookies: Optional[Dict[str, str]] = None,
        basic_auth: Optional[aiohttp.BasicAuth] = None
    ) -> Optional[str]:
        """Fetch URL with optional authentication (Bearer, Basic Auth, or Cookies)"""
        # Skip binary files
        skip_extensions = ['.pdf', '.zip', '.tar', '.gz', '.exe', '.dmg', '.pkg', '.deb', '.rpm', 
                          '.png', '.jpg', '.jpeg', '.gif', '.svg', '.ico', '.webp',
                          '.mp4', '.mp3', '.wav', '.avi', '.mov']
        if any(url.lower().endswith(ext) for ext in skip_extensions):
            print(f"Skipping binary file: {url}")
            self.scrape_stats["skipped"] += 1
            return None
        
        try:
            request_headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
                "Accept-Language": "en-US,en;q=0.5",
            }
            if headers:
                request_headers.update(headers)
            
            async with session.get(
                url, 
                timeout=aiohttp.ClientTimeout(total=30),
                headers=request_headers,
                cookies=cookies,
                auth=basic_auth,
                ssl=False  # Allow self-signed certs for internal wikis
            ) as response:
                if response.status == 200:
                    # Check content type to avoid binary files
                    content_type = response.headers.get('Content-Type', '')
                    if not any(t in content_type for t in ['text/', 'html', 'xml', 'json']):
                        print(f"Skipping non-text content: {url} ({content_type})")
                        self.scrape_stats["skipped"] += 1
                        return None
                    return await response.text()
                elif response.status == 401:
                    print(f"Authentication required for {url}")
                elif response.status == 403:
                    print(f"Access forbidden: {url}")
                else:
                    print(f"HTTP {response.status} for {url}")
        except Exception as e:
            print(f"Error fetching {url}: {e}")
            self.scrape_stats["errors"] += 1
        return None
    
    def detect_platform(self, html: str, url: str) -> str:
        """Auto-detect the documentation platform"""
        html_lower = html.lower()
        
        if "confluence" in html_lower or "atlassian" in html_lower:
            return "confluence"
        elif "sharepoint" in html_lower or "_layouts" in url:
            return "sharepoint"
        elif "gitbook" in html_lower or "gitbook.io" in url:
            return "gitbook"
        # Check sphinx BEFORE docusaurus since many sphinx sites use docsearch
        elif "sphinx" in html_lower or "sphinxdoc" in html_lower or "wy-nav-content" in html_lower or "rst-content" in html_lower:
            return "sphinx"
        elif "mkdocs" in html_lower or "material for mkdocs" in html_lower:
            return "mkdocs"
        elif "docusaurus" in html_lower:
            return "docusaurus"
        
        return "generic"
    
    def extract_text_from_html(self, html: str, url: str, platform: str = "auto") -> Document:
        """Extract text using platform-specific selectors"""
        soup = BeautifulSoup(html, 'html.parser')
        
        # Auto-detect platform if needed
        if platform == "auto":
            platform = self.detect_platform(html, url)
        
        selectors = PLATFORM_SELECTORS.get(platform, PLATFORM_SELECTORS["generic"])
        
        # Remove unwanted elements first
        for element in soup(['script', 'style', 'noscript', 'svg', 'iframe']):
            element.decompose()
        
        # Remove platform-specific noise
        for selector in selectors.get("remove", []):
            for element in soup.select(selector):
                element.decompose()
        
        # Find main content using platform-specific selectors
        main_content = None
        matched_selector = None
        for selector in selectors.get("content", []):
            main_content = soup.select_one(selector)
            if main_content:
                matched_selector = selector
                break
        
        # Fallback to generic detection
        if not main_content:
            main_content = soup.find('main') or soup.find('article') or soup.find('div', class_=re.compile(r'content|main|article', re.I))
            if main_content:
                matched_selector = f"fallback:{main_content.name}"
        
        if main_content:
            text = main_content.get_text(separator='\n', strip=True)
        else:
            text = soup.get_text(separator='\n', strip=True)
            matched_selector = "full-page"
        
        # Debug: log extraction details
        print(f"[Extract] Platform: {platform}, Selector: {matched_selector}, Raw text len: {len(text)}")
        
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        cleaned_text = '\n'.join(lines)
        
        title = None
        for title_selector in selectors.get("title", ["h1", "title"]):
            title_elem = soup.select_one(title_selector)
            if title_elem:
                title = title_elem.get_text(strip=True)
                break
        
        if not title:
            title_elem = soup.find('title')
            title = title_elem.get_text(strip=True) if title_elem else url
        
        return Document(
            page_content=cleaned_text,
            metadata={
                'source': url,
                'title': title,
                'type': 'html',
                'platform': platform
            }
        )
    
    def should_include_url(self, url: str, config: ScrapeConfig) -> bool:
        """Check if URL should be included based on patterns"""
        # Check exclude patterns first
        for pattern in config.exclude_patterns:
            if re.search(pattern, url):
                return False
        
        # If include patterns specified, URL must match at least one
        if config.url_patterns:
            for pattern in config.url_patterns:
                if re.search(pattern, url):
                    return True
            return False
        
        return True
    
    def extract_links(self, html: str, base_url: str, allowed_domains: List[str], config: Optional[ScrapeConfig] = None) -> List[str]:
        soup = BeautifulSoup(html, 'html.parser')
        links = []
        
        # Parse the base URL to get the current path context
        base_parsed = urlparse(base_url)
        
        # For versioned doc sites (like /en/stable/), extract the version prefix
        # to handle root-relative links that should stay within the version
        version_prefix = ""
        path_parts = base_parsed.path.split('/')
        # Check for patterns like /en/stable/, /en/v1.18/, /docs/v2/, etc.
        # path_parts for /en/stable/foo.html = ['', 'en', 'stable', 'foo.html']
        if len(path_parts) >= 4:
            # Check if this looks like a versioned docs path
            if path_parts[1] in ['en', 'docs', 'documentation']:
                if path_parts[2] in ['stable', 'latest', 'main', 'master'] or re.match(r'v?\d+\.?\d*', path_parts[2]):
                    version_prefix = '/' + path_parts[1] + '/' + path_parts[2]  # e.g., /en/stable
        
        for a_tag in soup.find_all('a', href=True):
            href = a_tag['href']
            
            # Skip anchors, javascript, mailto, etc.
            if href.startswith('#') or href.startswith('javascript:') or href.startswith('mailto:'):
                continue
            
            # Handle root-relative URLs that should stay within version prefix
            # Only modify if it's a root-relative link that doesn't already have the full version path
            if href.startswith('/') and not href.startswith('//') and version_prefix:
                # Check if href already contains the version prefix
                if version_prefix not in href:
                    # This is a link like /overview/ that should be /en/stable/overview/
                    href = version_prefix + href
                elif href.startswith('/' + path_parts[1] + '/') and not href.startswith(version_prefix):
                    # This is a link like /en/overview/ that should be /en/stable/overview/
                    # Replace /en/ with /en/stable/
                    href = href.replace('/' + path_parts[1] + '/', version_prefix + '/', 1)
            
            # Use urljoin which properly handles relative URLs
            full_url = urljoin(base_url, href)
            parsed = urlparse(full_url)
            
            if parsed.scheme not in ['http', 'https']:
                continue
            
            if any(domain in parsed.netloc for domain in allowed_domains):
                # Remove query string and fragment, keep path
                clean_url = f"{parsed.scheme}://{parsed.netloc}{parsed.path}"
                
                # Remove trailing slash for consistency (except for root)
                if clean_url.endswith('/') and len(parsed.path) > 1:
                    clean_url = clean_url.rstrip('/')
                
                # Apply URL pattern filters if config provided
                if config and not self.should_include_url(clean_url, config):
                    continue
                
                if clean_url not in self.visited_urls:
                    links.append(clean_url)
        
        return links
    
    async def scrape_url(
        self,
        url: str,
        recursive: bool = False,
        max_depth: int = 2,
        allowed_domains: Optional[List[str]] = None,
        job=None  # Optional job object for status tracking
    ) -> List[Document]:
        """Simple scrape method for backward compatibility"""
        config = ScrapeConfig(
            url=url,
            recursive=recursive,
            max_depth=max_depth,
            allowed_domains=allowed_domains or []
        )
        return await self.scrape_with_config(config, job=job)
    
    async def scrape_with_config(self, config: ScrapeConfig, job=None) -> List[Document]:
        """Advanced scraping with full configuration"""
        if not config.allowed_domains:
            parsed = urlparse(config.url)
            config.allowed_domains = [parsed.netloc]
        
        documents = []
        urls_to_process = [(config.url, 0)]
        
        # Update job with initial count
        if job:
            job.pages_found = 1
        
        # Prepare headers with auth if provided
        headers = dict(config.headers)
        if config.auth_token:
            headers["Authorization"] = f"Bearer {config.auth_token}"
        
        # Prepare Basic Auth if provided
        basic_auth = None
        if config.basic_auth_username and config.basic_auth_password:
            basic_auth = aiohttp.BasicAuth(config.basic_auth_username, config.basic_auth_password)
        
        # Merge cookies: explicit dict + parsed cookie string
        cookies = dict(config.cookies)
        if config.cookie_string:
            parsed_cookies = parse_cookie_string(config.cookie_string)
            cookies.update(parsed_cookies)
        
        connector = aiohttp.TCPConnector(ssl=False)  # Allow self-signed certs
        print(f"[Scrape] Starting crawl of {config.url}, max_pages={config.max_pages}, recursive={config.recursive}")
        async with aiohttp.ClientSession(connector=connector) as session:
            while urls_to_process and len(documents) < config.max_pages:
                # Check for cancellation
                if job and job.cancelled:
                    print("Crawl cancelled by user")
                    break
                
                current_url, depth = urls_to_process.pop(0)
                print(f"[Scrape] Processing: {current_url} (depth={depth})")
                
                if current_url in self.visited_urls:
                    print(f"[Scrape] Already visited: {current_url}")
                    continue
                
                # Check URL patterns
                if not self.should_include_url(current_url, config):
                    print(f"[Scrape] Excluded by pattern: {current_url}")
                    self.scrape_stats["skipped"] += 1
                    continue
                
                self.visited_urls.add(current_url)
                
                # Update job status
                if job:
                    job.current_page = current_url
                    job.pages_processed = self.scrape_stats["processed"]
                
                # Rate limiting
                if config.rate_limit > 0:
                    await asyncio.sleep(config.rate_limit)
                
                html = await self.fetch_url(session, current_url, headers, cookies, basic_auth)
                if not html:
                    print(f"[Scrape] No HTML returned for: {current_url}")
                    continue
                
                print(f"[Scrape] Got {len(html)} bytes from {current_url}")
                
                doc = self.extract_text_from_html(html, current_url, config.platform)
                content_len = len(doc.page_content) if doc.page_content else 0
                print(f"[Scrape] Extracted {content_len} chars from {current_url}")
                if doc.page_content and content_len > 100:
                    documents.append(doc)
                    self.scrape_stats["processed"] += 1
                    print(f"[{self.scrape_stats['processed']}/{config.max_pages}] {current_url} ({content_len} chars)")
                else:
                    print(f"[Scrape] SKIPPED (too short): {current_url} - only {content_len} chars")
                    
                    # Update job progress
                    if job:
                        job.pages_processed = self.scrape_stats["processed"]
                
                if config.recursive and depth < config.max_depth:
                    links = self.extract_links(html, current_url, config.allowed_domains, config)
                    for link in links:
                        if link not in self.visited_urls:
                            urls_to_process.append((link, depth + 1))
                    
                    # Update pages found count
                    if job:
                        job.pages_found = len(urls_to_process) + self.scrape_stats["processed"]
        
        print(f"\nScrape complete: {self.scrape_stats['processed']} pages, {self.scrape_stats['skipped']} skipped, {self.scrape_stats['errors']} errors")
        return documents
    
    def load_markdown_file(self, filepath: str) -> Document:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return Document(
            page_content=content,
            metadata={
                'source': filepath,
                'type': 'markdown',
                'filename': os.path.basename(filepath)
            }
        )
    
    def load_text_file(self, filepath: str) -> Document:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return Document(
            page_content=content,
            metadata={
                'source': filepath,
                'type': 'text',
                'filename': os.path.basename(filepath)
            }
        )
    
    def load_pdf_file(self, filepath: str) -> Document:
        """Load and extract text from a PDF file"""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(filepath)
            text_parts = []
            
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {page_num}]\n{page_text}")
            
            content = "\n\n".join(text_parts)
            
            # Extract metadata if available
            metadata = reader.metadata or {}
            title = metadata.get('/Title', '') or os.path.basename(filepath)
            
            return Document(
                page_content=content,
                metadata={
                    'source': filepath,
                    'type': 'pdf',
                    'filename': os.path.basename(filepath),
                    'title': title,
                    'pages': len(reader.pages)
                }
            )
        except Exception as e:
            print(f"Error loading PDF {filepath}: {e}")
            return Document(
                page_content="",
                metadata={'source': filepath, 'type': 'pdf', 'error': str(e)}
            )
    
    def load_docx_file(self, filepath: str) -> Document:
        """Load and extract text from a Word document"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(filepath)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            content = "\n\n".join(text_parts)
            
            # Try to get title from core properties
            title = os.path.basename(filepath)
            try:
                if doc.core_properties.title:
                    title = doc.core_properties.title
            except:
                pass
            
            return Document(
                page_content=content,
                metadata={
                    'source': filepath,
                    'type': 'docx',
                    'filename': os.path.basename(filepath),
                    'title': title
                }
            )
        except Exception as e:
            print(f"Error loading DOCX {filepath}: {e}")
            return Document(
                page_content="",
                metadata={'source': filepath, 'type': 'docx', 'error': str(e)}
            )
    
    def load_directory(self, directory: str, extensions: Optional[List[str]] = None) -> List[Document]:
        if extensions is None:
            extensions = ['.md', '.txt', '.rst', '.html', '.pdf', '.docx']
        
        documents = []
        
        for root, _, files in os.walk(directory):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in extensions:
                    filepath = os.path.join(root, file)
                    try:
                        if ext == '.md':
                            doc = self.load_markdown_file(filepath)
                        elif ext == '.pdf':
                            doc = self.load_pdf_file(filepath)
                        elif ext == '.docx':
                            doc = self.load_docx_file(filepath)
                        else:
                            doc = self.load_text_file(filepath)
                        if doc.page_content:  # Only add if content was extracted
                            documents.append(doc)
                            print(f"Loaded: {filepath}")
                    except Exception as e:
                        print(f"Error loading {filepath}: {e}")
        
        return documents
    
    def reset_visited(self):
        self.visited_urls.clear()
        self.scrape_stats = {"processed": 0, "skipped": 0, "errors": 0}


# ============== Preset Configurations (YAML-driven) ==============

def _build_scrape_config_from_preset(preset) -> ScrapeConfig:
    """Convert a domain_config ScrapePresetConfig into a ScrapeConfig."""
    return ScrapeConfig(
        url=preset.url,
        recursive=preset.recursive,
        max_depth=preset.max_depth,
        max_pages=preset.max_pages,
        allowed_domains=list(preset.allowed_domains),
        url_patterns=list(preset.url_patterns),
        exclude_patterns=list(preset.exclude_patterns),
        rate_limit=preset.rate_limit,
        platform=preset.platform,
    )


def _load_scrape_presets() -> dict:
    """Build the preset registry from the domain YAML config."""
    from app.domain_config import get_domain_config
    dc = get_domain_config()
    return {
        name: (lambda p=preset: _build_scrape_config_from_preset(p))
        for name, preset in dc.scrape_presets.items()
    }


def get_confluence_config(base_url: str, space_key: str, auth_token: Optional[str] = None) -> ScrapeConfig:
    """Preset for Confluence wiki scraping"""
    return ScrapeConfig(
        url=f"{base_url}/wiki/spaces/{space_key}",
        recursive=True,
        max_depth=5,
        max_pages=500,
        allowed_domains=[urlparse(base_url).netloc],
        url_patterns=[rf"/wiki/spaces/{space_key}/"],
        exclude_patterns=[
            r"/wiki/spaces/.*/history",
            r"/wiki/spaces/.*/attachments",
            r"action=edit",
        ],
        auth_token=auth_token,
        rate_limit=0.5,
        platform="confluence"
    )


def get_generic_wiki_config(base_url: str, auth_token: Optional[str] = None) -> ScrapeConfig:
    """Generic configuration for internal wikis"""
    parsed = urlparse(base_url)
    return ScrapeConfig(
        url=base_url,
        recursive=True,
        max_depth=4,
        max_pages=500,
        allowed_domains=[parsed.netloc],
        exclude_patterns=[
            r"/edit",
            r"/history",
            r"/diff",
            r"action=",
            r"\?.*edit",
        ],
        auth_token=auth_token,
        rate_limit=0.5,
        platform="auto"
    )


def get_scrape_presets() -> dict:
    """Return the current preset registry (lazy-loaded from YAML)."""
    return _load_scrape_presets()


# Keep a module-level alias for backward compatibility with imports
SCRAPE_PRESETS = None  # Populated lazily; use get_scrape_presets()


def _ensure_presets():
    global SCRAPE_PRESETS
    SCRAPE_PRESETS = get_scrape_presets()
    return SCRAPE_PRESETS


async def scrape_bulk_docs() -> List[Document]:
    """Scrape all presets listed in bulk_scrape_presets and load local_directories from domain YAML."""
    from app.domain_config import get_domain_config
    dc = get_domain_config()
    presets = get_scrape_presets()

    loader = DocumentLoader()
    all_documents = []

    # Load local directories first (fast, no network)
    for name, dir_cfg in dc.local_directories.items():
        if not os.path.isdir(dir_cfg.path):
            print(f"[BulkLoad] Skipping '{name}': directory not found at {dir_cfg.path}")
            continue
        print(f"\nLoading local directory '{name}' from {dir_cfg.path}...")
        docs = loader.load_directory(dir_cfg.path, extensions=dir_cfg.extensions)
        all_documents.extend(docs)
        print(f"[BulkLoad] Loaded {len(docs)} documents from '{name}'")

    # Then scrape web presets
    bulk_names = dc.bulk_scrape_presets or list(presets.keys())
    for name in bulk_names:
        if name not in presets:
            print(f"[BulkScrape] Skipping unknown preset: {name}")
            continue
        print(f"\nScraping '{name}' documentation...")
        config = presets[name]()
        docs = await loader.scrape_with_config(config)
        all_documents.extend(docs)
        loader.reset_visited()

    print(f"\nTotal documents loaded: {len(all_documents)}")
    return all_documents


async def scrape_with_preset(preset_name: str) -> List[Document]:
    """Scrape using a named preset configuration"""
    presets = get_scrape_presets()
    if preset_name not in presets:
        raise ValueError(f"Unknown preset: {preset_name}. Available: {list(presets.keys())}")

    config = presets[preset_name]()
    loader = DocumentLoader()
    return await loader.scrape_with_config(config)
